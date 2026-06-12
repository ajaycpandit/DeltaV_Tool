from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import io, re, os, openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

app = Flask(__name__)
CORS(app, expose_headers=['X-Sheet-Names'])

# ═══════════════════════════════════════════════════════════════════════════════
# SHARED UTILITIES
# ═══════════════════════════════════════════════════════════════════════════════

def decode_fhx(raw):
    if raw[:2] == b'\xff\xfe': return raw.decode('utf-16-le', errors='replace').lstrip('\ufeff')
    if raw[:2] == b'\xfe\xff': return raw.decode('utf-16-be', errors='replace').lstrip('\ufeff')
    return raw.decode('utf-8', errors='replace')

def extract_block(text, start):
    depth, i = 0, start
    while i < len(text):
        if text[i] == '{': depth += 1
        elif text[i] == '}':
            depth -= 1
            if depth == 0: return text[start:i+1]
        i += 1
    return ''

def parse_actions(sc_text):
    actions = []
    for m in re.finditer(r'ACTION\s+NAME="([^"]+)"\s*\{', sc_text):
        b = extract_block(sc_text, m.end()-1)
        g = lambda p: (re.search(p, b) or type('x',(),{'group':lambda s,n:''})()).group(1)
        expr_m = re.search(r'EXPRESSION="(.*?)"(?=\s+(?:DELAY|CONFIRM)|\s*\})', b, re.DOTALL)
        actions.append({
            'action':             m.group(1),
            'description':        g('DESCRIPTION="([^"]+)"'),
            'type':               g('ACTION_TYPE=(\\S+)'),
            'qualifier':          g('QUALIFIER=(\\S+)'),
            'expression':         expr_m.group(1).strip() if expr_m else '',
            'delay_time':         g('DELAY_TIME=(\\S+)'),
            'delay_expression':   g('DELAY_EXPRESSION="([^"]+)"'),
            'confirm_expression': g('CONFIRM_EXPRESSION="([^"]+)"'),
            'confirm_timeout':    g('CONFIRM_TIME_OUT=(\\S+)'),
        })
    return actions

def parse_sfc(sfc_block):
    steps, trans = {}, {}
    for sm in re.finditer(r'STEP\s+NAME="([^"]+)"\s*\{', sfc_block):
        sb = extract_block(sfc_block, sm.end()-1)
        rc = re.search(r'RECTANGLE\s*=\s*\{\s*X=(\d+)\s*Y=(\d+)', sb)
        steps[sm.group(1)] = {
            'x': int(rc.group(1)) if rc else 0,
            'y': int(rc.group(2)) if rc else 0,
            'actions': parse_actions(sb)
        }
    for tm in re.finditer(r'TRANSITION\s+NAME="([^"]+)"\s*\{', sfc_block):
        tb  = extract_block(sfc_block, tm.end()-1)
        pos = re.search(r'POSITION\s*=\s*\{\s*X=(\d+)\s*Y=(\d+)', tb)
        td  = re.search(r'DESCRIPTION="([^"]+)"', tb)
        te  = re.search(r'EXPRESSION="([^"]*)"', tb, re.DOTALL)
        tt  = re.search(r'TERMINATION=(\w+)', tb)
        trans[tm.group(1)] = {
            'x': int(pos.group(1)) if pos else 0,
            'y': int(pos.group(2)) if pos else 0,
            'description': td.group(1).strip() if td else '',
            'expression':  te.group(1).strip() if te else '',
            'termination': tt.group(1) if tt else 'F',
        }
    s2t = {s: [] for s in steps}
    for tn, tp in trans.items():
        cands = [(tp['y']-v['y'], sn) for sn,v in steps.items()
                 if v['y'] < tp['y'] and abs(v['x']-tp['x']) < 150]
        if not cands:
            cands = [(tp['y']-v['y'], sn) for sn,v in steps.items() if v['y'] < tp['y']]
        if cands:
            s2t[sorted(cands)[0][1]].append(tn)
    return {
        'ordered_steps': sorted(steps.items(), key=lambda x: x[1]['y']),
        'transitions':   trans,
        'step_to_trans': s2t,
    }

# ═══════════════════════════════════════════════════════════════════════════════
# SHARED EXCEL STYLES
# ═══════════════════════════════════════════════════════════════════════════════

NAVY    = PatternFill('solid', start_color='1E3A5F')
BLUE_H  = PatternFill('solid', start_color='2C4A72')
BLUE_S  = PatternFill('solid', start_color='3B6EA8')
TEAL_S  = PatternFill('solid', start_color='2B6070')
GREEN_S = PatternFill('solid', start_color='3D6B4F')
ORANGE_H= PatternFill('solid', start_color='6B5230')
OPEN_F  = PatternFill('solid', start_color='D4EDDA')
CLOSE_F = PatternFill('solid', start_color='F5DDDD')
DC_F    = PatternFill('solid', start_color='FDFAE8')
ALT     = PatternFill('solid', start_color='EDF2F8')
ALT_G   = PatternFill('solid', start_color='EDF5EC')
ALTG2   = PatternFill('solid', start_color='F4F9F3')
ALT_ROW = PatternFill('solid', start_color='F5F5F3')
WHITE   = PatternFill('solid', start_color='FFFFFF')
DIS_F   = PatternFill('solid', start_color='F0F0EE')
THIN    = Side(style='thin', color='AAAAAA')
BORD    = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
NCOLS   = 10
COL_W   = [14, 22, 32, 10, 10, 58, 10, 38, 34, 14]

def wf(bold=False, sz=10, color='000000'):
    return Font(name='Calibri', bold=bold, size=sz, color=color)
def wa(h='left', wrap=False):
    return Alignment(horizontal=h, vertical='top', wrap_text=wrap)
def sc(ws, r, c, val='', bold=False, sz=10, fc='000000', fill=None,
       h='left', wrap=False, merge_to=None):
    cell = ws.cell(row=r, column=c, value=str(val) if val is not None else '')
    cell.font = wf(bold, sz, fc); cell.fill = fill or WHITE
    cell.alignment = wa(h, wrap); cell.border = BORD
    if merge_to:
        ws.merge_cells(start_row=r, start_column=c, end_row=r, end_column=merge_to)
    return cell

def write_sfc_sheet(wb, label, title_fill, step_fill, data, opts, fname_desc):
    ws = wb.create_sheet(title=label[:31])
    for ci, w in enumerate(COL_W, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.freeze_panes = 'A3'

    # Title
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=NCOLS)
    t = ws.cell(row=1, column=1, value=f"  {fname_desc}")
    t.font = wf(True,12,'FFFFFF'); t.fill = NAVY
    t.alignment = Alignment(horizontal='left', vertical='center'); t.border = BORD
    ws.row_dimensions[1].height = 24

    # Column headers
    hdrs = ['Row Type','Step / Transition','Description','Action','Qualifier',
            'Expression' if opts.get('expressions',True) else '(hidden)',
            'Delay','Delay Expression','Confirm Expression','Confirm Timeout']
    for ci, h in enumerate(hdrs, 1):
        sc(ws, 2, ci, h, bold=True, fc='FFFFFF', fill=BLUE_H, h='center')
    ws.row_dimensions[2].height = 17

    row = 3
    for si, (sn, sd) in enumerate(data['ordered_steps']):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=NCOLS)
        sh = ws.cell(row=row, column=1,
                     value=f"  STEP {si+1}:  {sn}   "
                           f"({len(sd['actions'])} action{'s' if len(sd['actions'])!=1 else ''})")
        sh.font = wf(True,10,'FFFFFF'); sh.fill = step_fill
        sh.alignment = Alignment(horizontal='left', vertical='center'); sh.border = BORD
        ws.row_dimensions[row].height = 17; row += 1

        for ai, a in enumerate(sd['actions']):
            f = ALT if ai%2==0 else WHITE
            ie = opts.get('expressions', True)
            sc(ws,row,1, 'ACTION',                              bold=True, fill=f, h='center')
            sc(ws,row,2, sn,                                    fill=f)
            sc(ws,row,3, a['description'],                      fill=f, wrap=True)
            sc(ws,row,4, a['action'],                           fill=f, h='center')
            sc(ws,row,5, a['qualifier'],                        fill=f, h='center')
            sc(ws,row,6, a['expression'] if ie else '—',        fill=f, wrap=True)
            sc(ws,row,7, a['delay_time'],                       fill=f, h='center')
            sc(ws,row,8, a['delay_expression'],                 fill=f, wrap=True)
            sc(ws,row,9, a['confirm_expression'],               fill=f, wrap=True)
            sc(ws,row,10,a['confirm_timeout'],                  fill=f, h='center')
            ws.row_dimensions[row].height = max(15, min(75,
                15*max(1, len(a.get('expression',''))//55+1)))
            row += 1

        if opts.get('transitions', True):
            tl = data['step_to_trans'].get(sn, [])
            if tl:
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=NCOLS)
                th = ws.cell(row=row, column=1,
                             value=f"  ↓  TRANSITION{'S' if len(tl)>1 else ''} "
                                   f"FROM  {sn}  ({len(tl)})")
                th.font = wf(True,9,'FFFFFF'); th.fill = GREEN_S
                th.alignment = Alignment(horizontal='left', vertical='center'); th.border = BORD
                ws.row_dimensions[row].height = 15; row += 1

                for ti, tn in enumerate(tl):
                    tr = data['transitions'][tn]
                    f  = ALT_G if ti%2==0 else ALTG2
                    sc(ws,row,1,'TRANSITION', bold=True, fill=f, h='center')
                    sc(ws,row,2,tn,           bold=True, fill=f)
                    sc(ws,row,3,tr['description'], fill=f, wrap=True)
                    sc(ws,row,4,'⏹ END' if tr['termination']=='T' else '→ NEXT',
                       fill=f, h='center')
                    sc(ws,row,5,'',fill=f)
                    ws.merge_cells(start_row=row, start_column=6, end_row=row, end_column=NCOLS)
                    ec = ws.cell(row=row, column=6,
                                 value=tr['expression'] if opts.get('expressions',True) else '—')
                    ec.font = wf(); ec.fill = f
                    ec.alignment = wa('left',True); ec.border = BORD
                    ws.row_dimensions[row].height = max(15, min(60,
                        15*max(1, len(tr.get('expression',''))//80+1)))
                    row += 1

        for ci in range(1, NCOLS+1): sc(ws, row, ci, '', fill=WHITE)
        ws.row_dimensions[row].height = 6; row += 1

    return ws

# ═══════════════════════════════════════════════════════════════════════════════
# PHASE FHX PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def build_instance_map(text):
    mapping = {}
    for m in re.finditer(r'FUNCTION_BLOCK\s+NAME="([^"]+)"\s+DEFINITION="([^"]+)"', text):
        inst, defn = m.group(1), m.group(2)
        if defn not in mapping:
            mapping[defn] = inst
    return mapping

def clean_label(raw):
    s = raw.strip().upper()
    s = re.sub(r'[\-_/\\]+', ' ', s)
    for noise in [' LOGIC',' PHASE',' SEQUENCE',' SUB',' FUNCTION']:
        s = s.replace(noise, '')
    s = re.sub(r'\s+', '_', s.strip())[:28]
    return re.sub(r'[\/\\\?\*\[\]\:\']', '_', s) or raw[:28]

def derive_phase_label(fb_name, instance_name, description, used_labels):
    if instance_name and instance_name.upper() not in ('','NONE'):
        candidate = clean_label(instance_name)
    else:
        desc_up = description.upper()
        _KW = [
            ('ABORT','ABORT_STOP'),('STOP','ABORT_STOP'),
            ('RUNNING','RUNNING'),('RUN ','RUNNING'),
            ('HOLDING','HOLD'),('HOLD','HOLD'),
            ('RESTART','RESTART'),('CONDITION','CONDITION'),
            ('PROMPT','PROMPT'),('INIT','INIT'),
            ('START','START'),('COMPLETE','COMPLETE'),
            ('FAIL','FAIL'),('PAUSE','PAUSE'),
        ]
        candidate = None
        for kw, lbl in _KW:
            if kw in desc_up:
                candidate = lbl; break
        if not candidate:
            words = re.sub(r'[^A-Z0-9\s]', ' ', desc_up).split()
            noise = {'LOGIC','PHASE','THE','FOR','AND','OR','A','AN','OF','FUNCTION','BLOCK','STEP','SUB','SEQUENCE'}
            words = [w for w in words if w not in noise and len(w)>1]
            candidate = '_'.join(words[-2:]) if words else fb_name[:20]
    label = re.sub(r'[\/\\\?\*\[\]\:\']', '_', candidate)[:28]
    base, n = label, 2
    while label in used_labels:
        label = f"{base[:25]}_{n}"; n += 1
    used_labels.add(label)
    return label

def parse_phase_fhx(text):
    instance_map = build_instance_map(text)
    blocks = {}
    for m in re.finditer(r'FUNCTION_BLOCK_DEFINITION\s+NAME="([^"]+)"[^\{]*\{', text):
        name  = m.group(1)
        block = extract_block(text, m.end()-1)
        desc  = re.search(r'DESCRIPTION="([^"]+)"', block)
        sfc_m = re.search(r'SFC_ALGORITHM\s*\{', block)
        sfc   = extract_block(block, sfc_m.end()-1) if sfc_m else ''
        sfc_data = parse_sfc(sfc)
        blocks[name] = {
            'description':   desc.group(1) if desc else '',
            'instance_name': instance_map.get(name, ''),
            **sfc_data,
        }
    return blocks

def build_phase_excel(blocks, fname, opts):
    wb = openpyxl.Workbook()
    used_labels = set()
    block_labels = {
        fb: derive_phase_label(fb, d.get('instance_name',''), d.get('description',''), used_labels)
        for fb, d in blocks.items()
    }

    if opts.get('summary', True):
        ws = wb.active; ws.title = 'SUMMARY'
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=8)
        t = ws.cell(row=1, column=1, value=f"  {fname}  —  DeltaV Phase Export Logic Summary")
        t.font = wf(True,13,'FFFFFF'); t.fill = NAVY
        t.alignment = Alignment(horizontal='left', vertical='center'); t.border = BORD
        ws.row_dimensions[1].height = 28
        for ci, h in enumerate(['Logic Block','FB Name','Description','Steps',
                                 'Actions','Transitions','Orphan Trans.','Sheet Tab'], 1):
            sc(ws, 2, ci, h, bold=True, fc='FFFFFF', fill=BLUE_H, h='center')
        ws.row_dimensions[2].height = 17
        for ci, w in enumerate([18,30,50,8,10,13,14,18], 1):
            ws.column_dimensions[get_column_letter(ci)].width = w
        ws.freeze_panes = 'A3'
        for i, (fb, data) in enumerate(blocks.items()):
            lbl    = block_labels[fb]
            acts   = sum(len(s[1]['actions']) for s in data['ordered_steps'])
            mapped = sum(len(v) for v in data['step_to_trans'].values())
            orphan = len(data['transitions']) - mapped
            f      = ALT if i%2==0 else WHITE
            r      = i+3
            sc(ws,r,1,lbl,bold=True,fill=f); sc(ws,r,2,fb,fill=f)
            sc(ws,r,3,data['description'],fill=f,wrap=True)
            sc(ws,r,4,len(data['ordered_steps']),fill=f,h='center')
            sc(ws,r,5,acts,fill=f,h='center')
            sc(ws,r,6,len(data['transitions']),fill=f,h='center')
            sc(ws,r,7,orphan or '-',fill=f,h='center')
            sc(ws,r,8,lbl,fill=f)
            ws.row_dimensions[r].height = 16
    else:
        wb.active.title = '_temp'

    for fb, data in blocks.items():
        lbl = block_labels[fb]
        write_sfc_sheet(wb, lbl, NAVY, BLUE_S, data, opts,
                        f"Phase Logic: {lbl}   |   {data['description']}")

    if '_temp' in wb.sheetnames and len(wb.sheetnames) > 1:
        del wb['_temp']
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════════
# COMMAND DRIVEN EM PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def parse_cdem_fhx(text):
    results = []
    for mc_m in re.finditer(r'MODULE_CLASS\s+NAME="([^"]+)"[^\{]*\{', text):
        em_name  = mc_m.group(1)
        mc_block = extract_block(text, mc_m.end()-1)
        if 'COMMAND_DRIVEN_ALGORITHM' not in mc_block:
            continue
        em_desc_m = re.search(r'DESCRIPTION="([^"]+)"', mc_block)
        em_desc   = em_desc_m.group(1) if em_desc_m else em_name

        cda_m = re.search(r'COMMAND_DRIVEN_ALGORITHM\s*\{', mc_block)
        cda   = extract_block(mc_block, cda_m.end()-1)

        cmd_index_to_name = {}
        for tm in re.finditer(r'TRANSITION\s+NAME="T_IN_(\d+)"\s*\{', cda):
            idx = int(tm.group(1))
            tb  = extract_block(cda, tm.end()-1)
            expr = re.search(r'EXPRESSION="([^"]+)"', tb)
            if expr:
                nm = re.search(r"'[^']+:([^']+)'", expr.group(1))
                if nm: cmd_index_to_name[idx] = nm.group(1).strip()

        cmd_inst_to_def = {}
        for fb_m in re.finditer(
            r'FUNCTION_BLOCK\s+NAME="(COMMAND_\d+)"\s+DEFINITION="([^"]+)"', mc_block):
            idx = int(re.search(r'(\d+)$', fb_m.group(1)).group(1))
            cmd_inst_to_def[idx] = (fb_m.group(1), fb_m.group(2))

        for idx in sorted(cmd_inst_to_def.keys()):
            inst, defn = cmd_inst_to_def[idx]
            cmd_name   = cmd_index_to_name.get(idx, inst)
            fb_m = re.search(
                r'FUNCTION_BLOCK_DEFINITION\s+NAME="' + re.escape(defn) + r'"[^\{]*\{', text)
            if not fb_m: continue
            fb_block = extract_block(text, fb_m.end()-1)
            fb_desc  = re.search(r'DESCRIPTION="([^"]+)"', fb_block)
            sfc_m    = re.search(r'SFC_ALGORITHM\s*\{', fb_block)
            if not sfc_m: continue
            sfc_data = parse_sfc(extract_block(fb_block, sfc_m.end()-1))
            results.append({
                'em_name':        em_name,
                'em_description': em_desc,
                'command_name':   cmd_name,
                'command_index':  idx,
                'fb_definition':  defn,
                'fb_description': fb_desc.group(1) if fb_desc else '',
                **sfc_data,
            })
    return results

def build_cdem_excel(commands, fname, opts):
    wb = openpyxl.Workbook()

    if opts.get('summary', True):
        ws = wb.active; ws.title = 'SUMMARY'
        em_name = commands[0]['em_name'] if commands else fname
        em_desc = commands[0]['em_description'] if commands else ''
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=8)
        t = ws.cell(row=1, column=1, value=f"  {em_name}  —  Command Driven EM Logic Summary")
        t.font = wf(True,13,'FFFFFF'); t.fill = NAVY
        t.alignment = Alignment(horizontal='left', vertical='center'); t.border = BORD
        ws.row_dimensions[1].height = 28
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=8)
        s = ws.cell(row=2, column=1, value=f"  {em_desc}")
        s.font = wf(False,10,'CCCCCC'); s.fill = NAVY
        s.alignment = Alignment(horizontal='left', vertical='center'); s.border = BORD
        ws.row_dimensions[2].height = 16
        for ci, h in enumerate(['#','Command','Description','Steps',
                                 'Actions','Transitions','Sheet Tab','FB Definition'], 1):
            sc(ws, 3, ci, h, bold=True, fc='FFFFFF', fill=BLUE_H, h='center')
        ws.row_dimensions[3].height = 17
        for ci, w in enumerate([5,18,42,7,10,12,16,30], 1):
            ws.column_dimensions[get_column_letter(ci)].width = w
        ws.freeze_panes = 'A4'
        for i, cmd in enumerate(commands):
            acts = sum(len(s[1]['actions']) for s in cmd['ordered_steps'])
            f    = ALT if i%2==0 else WHITE; r = i+4
            sc(ws,r,1,cmd['command_index'],     fill=f,h='center')
            sc(ws,r,2,cmd['command_name'],       fill=f,bold=True)
            sc(ws,r,3,cmd['fb_description'],     fill=f,wrap=True)
            sc(ws,r,4,len(cmd['ordered_steps']), fill=f,h='center')
            sc(ws,r,5,acts,                      fill=f,h='center')
            sc(ws,r,6,len(cmd['transitions']),   fill=f,h='center')
            sc(ws,r,7,cmd['command_name'][:31],  fill=f)
            sc(ws,r,8,cmd['fb_definition'],      fill=f)
            ws.row_dimensions[r].height = 16
    else:
        wb.active.title = '_temp'

    used_tabs = set()
    for cmd in commands:
        base = re.sub(r'[\/\\\?\*\[\]\:\s]', '_', cmd['command_name'])[:28]
        tab  = base; n = 2
        while tab in used_tabs: tab = f"{base[:25]}_{n}"; n += 1
        used_tabs.add(tab)
        write_sfc_sheet(wb, tab, NAVY, TEAL_S, cmd, opts,
                        f"EM Command: {cmd['command_name']}   |   "
                        f"{cmd['em_name']}  —  {cmd['fb_description']}")

    if '_temp' in wb.sheetnames and len(wb.sheetnames) > 1:
        del wb['_temp']
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════════
# STATE DRIVEN EM PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def parse_sdem_fhx(text):
    results = []
    for mc_m in re.finditer(r'MODULE_CLASS\s+NAME="([^"]+)"[^\{]*\{', text):
        em_name  = mc_m.group(1)
        mc_block = extract_block(text, mc_m.end()-1)
        if 'STATE_DRIVEN_ALGORITHM' not in mc_block:
            continue
        em_desc_m = re.search(r'DESCRIPTION="([^"]+)"', mc_block)
        em_desc   = em_desc_m.group(1) if em_desc_m else em_name

        sda_m = re.search(r'STATE_DRIVEN_ALGORITHM\s*\{', mc_block)
        sda   = extract_block(mc_block, sda_m.end()-1)
        sfc_m = re.search(r'SFC_ALGORITHM\s*\{', sda)
        if not sfc_m: continue
        sfc = extract_block(sda, sfc_m.end()-1)

        index_to_name = {}
        for tm in re.finditer(r'TRANSITION\s+NAME="T_IN_(\d+)"\s*\{', sfc):
            idx  = int(tm.group(1))
            tb   = extract_block(sfc, tm.end()-1)
            expr = re.search(r'EXPRESSION="([^"]+)"', tb)
            if expr:
                nm = re.search(r"'[^']+:([^']+)'", expr.group(1))
                if nm: index_to_name[idx] = nm.group(1).strip()

        state_def = None
        for fb_m in re.finditer(
            r'FUNCTION_BLOCK\s+NAME="STATE_\d+"\s+DEFINITION="([^"]+)"', mc_block):
            state_def = fb_m.group(1); break
        if not state_def: continue

        fb_def_m = re.search(
            r'FUNCTION_BLOCK_DEFINITION\s+NAME="' + re.escape(state_def) + r'"[^\{]*\{', text)
        if not fb_def_m: continue
        fb_def_block = extract_block(text, fb_def_m.end()-1)

        devices, default_vals, default_dc = [], {}, {}
        for am in re.finditer(r'ATTRIBUTE\s+NAME="([^"]+)"[^\{]*\{', fb_def_block):
            ab = extract_block(fb_def_block, am.end()-1)
            if re.search(r'GROUP="I/O"', ab):
                devices.append(am.group(1))
        for ai_m in re.finditer(r'ATTRIBUTE_INSTANCE\s+NAME="([^"]+)"\s*\{', fb_def_block):
            ai_b  = extract_block(fb_def_block, ai_m.end()-1)
            dname = ai_m.group(1)
            val_m = re.search(r'STRING_VALUE="([^"]+)"', ai_b)
            cv_m  = re.search(r'CV=([TF\d\.]+)', ai_b)
            dc_m  = re.search(r'SDA_DONT_CARE=([TF])', ai_b)
            default_vals[dname] = val_m.group(1) if val_m else (cv_m.group(1) if cv_m else '')
            default_dc[dname]   = dc_m.group(1)=='T' if dc_m else False

        state_overrides, enabled_overrides = {}, {}
        for ai_m in re.finditer(
            r'ATTRIBUTE_INSTANCE\s+NAME="STATE_(\d+)/([^"]+)"\s*\{', mc_block):
            idx, dname = int(ai_m.group(1)), ai_m.group(2)
            ai_b  = extract_block(mc_block, ai_m.end()-1)
            val_m = re.search(r'STRING_VALUE="([^"]+)"', ai_b)
            cv_m  = re.search(r'CV=([TF\d\.]+)', ai_b)
            dc_m  = re.search(r'SDA_DONT_CARE=([TF])', ai_b)
            val   = val_m.group(1) if val_m else (cv_m.group(1) if cv_m else '')
            dc    = dc_m.group(1)=='T' if dc_m else False
            if dname == 'ENABLED':
                enabled_overrides[idx] = (cv_m.group(1) != 'F') if cv_m else True
            else:
                state_overrides.setdefault(idx, {})[dname] = (val, dc)

        states = []
        for idx in sorted(index_to_name.keys()):
            ov  = state_overrides.get(idx, {})
            row_vals, row_dc = {}, {}
            for dev in devices:
                if dev in ov: row_vals[dev], row_dc[dev] = ov[dev]
                else: row_vals[dev] = default_vals.get(dev,''); row_dc[dev] = default_dc.get(dev,False)
            states.append({
                'index':      idx,
                'state_name': index_to_name[idx],
                'enabled':    enabled_overrides.get(idx, True),
                'values':     row_vals,
                'dont_care':  row_dc,
            })

        results.append({
            'em_name':      em_name,
            'em_description': em_desc,
            'devices':      devices,
            'default_vals': default_vals,
            'states':       states,
        })
    return results

def build_sdem_excel(em_list, fname, opts):
    wb = openpyxl.Workbook(); first = True

    for em in em_list:
        devices, states = em['devices'], em['states']
        em_name, em_desc = em['em_name'], em['em_description']
        tab = re.sub(r'[\/\\\?\*\[\]\:\s]', '_', em_name)[:31]
        ws  = wb.active if first else wb.create_sheet(title=tab)
        if first: ws.title = tab; first = False
        NCOLS_T = 3 + len(devices) + 1

        # Title
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=NCOLS_T)
        t = ws.cell(row=1, column=1,
                    value=f"  {em_name}  —  State Driven EM  |  {em_desc}")
        t.font = wf(True,13,'FFFFFF'); t.fill = NAVY
        t.alignment = Alignment(horizontal='left', vertical='center'); t.border = BORD
        ws.row_dimensions[1].height = 28

        # Device header row
        for ci in range(1, 4):
            sc(ws, 2, ci, '', fill=BLUE_H)
        for ci, dev in enumerate(devices, 4):
            sc(ws, 2, ci, dev, bold=True, fc='FFFFFF', fill=ORANGE_H, sz=11, h='center')
        sc(ws, 2, 4+len(devices), 'Min Time', bold=True, fc='FFFFFF', fill=BLUE_H, h='center')
        ws.row_dimensions[2].height = 20

        # Column headers
        for ci, h in enumerate(['#','State Name','Enabled'], 1):
            sc(ws, 3, ci, h, bold=True, fc='FFFFFF', fill=BLUE_H, h='center')
        for ci, dev in enumerate(devices, 4):
            sc(ws, 3, ci, dev, bold=True, fc='FFFFFF', fill=BLUE_H, h='center')
        sc(ws, 3, 4+len(devices), 'CFM Min Time', bold=True, fc='FFFFFF', fill=BLUE_H, h='center')
        ws.row_dimensions[3].height = 17
        ws.freeze_panes = 'A4'

        # State rows
        for ri, state in enumerate(states):
            r   = ri+4
            bg  = ALT_ROW if ri%2==0 else WHITE
            dis = not state['enabled']
            sc(ws,r,1, state['index'],      fill=DIS_F if dis else bg, h='center')
            sc(ws,r,2, state['state_name'], bold=True, fill=DIS_F if dis else bg,
               h='left', fc='888888' if dis else '000000')
            sc(ws,r,3, 'No' if dis else 'Yes',
               fill=DIS_F if dis else bg,
               fc='888888' if dis else '006100', h='center')
            for ci, dev in enumerate(devices, 4):
                val, dc = state['values'].get(dev,''), state['dont_care'].get(dev,False)
                if dis:     fill_c = DIS_F
                elif dc:    fill_c = DC_F
                elif val.upper()=='OPEN':  fill_c = OPEN_F
                elif val.upper()=='CLOSE': fill_c = CLOSE_F
                else:       fill_c = bg
                display = '—' if dc else val
                sc(ws,r,ci, display, fill=fill_c, h='center',
                   fc='888888' if dis or dc else '000000',
                   bold=(val.upper()=='OPEN' and not dc and not dis))
            sc(ws,r, 4+len(devices), '', fill=DIS_F if dis else bg, h='center')
            ws.row_dimensions[r].height = 16

        # Legend
        lr = len(states)+5
        ws.merge_cells(start_row=lr, start_column=1, end_row=lr, end_column=NCOLS_T)
        lh = ws.cell(row=lr, column=1, value='  LEGEND')
        lh.font = wf(True,10,'FFFFFF'); lh.fill = BLUE_H
        lh.alignment = Alignment(horizontal='left', vertical='center'); lh.border = BORD
        ws.row_dimensions[lr].height = 16
        for li, (fill_c, lbl) in enumerate([
            (OPEN_F,  'OPEN  —  device commanded open / active'),
            (CLOSE_F, 'CLOSE  —  device commanded closed / inactive'),
            (DC_F,    '—  (Don\'t Care)  —  device not controlled in this state'),
            (DIS_F,   'State disabled / not used'),
        ]):
            rr = lr+1+li
            sc(ws, rr, 1, '', fill=fill_c)
            ws.merge_cells(start_row=rr, start_column=2, end_row=rr, end_column=NCOLS_T)
            lc = ws.cell(row=rr, column=2, value=f'  {lbl}')
            lc.font = wf(False,10); lc.fill = WHITE
            lc.alignment = Alignment(horizontal='left', vertical='center'); lc.border = BORD
            ws.row_dimensions[rr].height = 15

        # Column widths
        ws.column_dimensions['A'].width = 6
        ws.column_dimensions['B'].width = 36
        ws.column_dimensions['C'].width = 9
        for ci in range(4, 4+len(devices)):
            ws.column_dimensions[get_column_letter(ci)].width = 14
        ws.column_dimensions[get_column_letter(4+len(devices))].width = 14

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════════
# AUTO-DETECT
# ═══════════════════════════════════════════════════════════════════════════════

def detect_fhx_type(text):
    """Return 'phase', 'em_cd', or 'em_sd' by inspecting the FHX structure."""
    if 'STATE_DRIVEN_ALGORITHM' in text:
        return 'em_sd'
    if 'COMMAND_DRIVEN_ALGORITHM' in text:
        return 'em_cd'
    return 'phase'

# ═══════════════════════════════════════════════════════════════════════════════
# SFC DIAGRAM BUILDER
# Generates a cell-grid SFC diagram in Excel with clickable step/transition
# shapes that hyperlink to their detail rows on a sibling sheet.
# ═══════════════════════════════════════════════════════════════════════════════

SCALE_D  = 0.09   # DeltaV coord units → grid cells
STEP_W_D = 18     # step box width  (grid cells)
STEP_H_D = 4      # step box height
TRAN_W_D = 10     # transition width
TRAN_H_D = 2      # transition height

def _to_grid(x, y, x_min, y_min):
    col = max(2, round((x - x_min) * SCALE_D) + 3)
    row = max(2, round((y - y_min) * SCALE_D) + 2)
    return row, col

def _draw_step_cell(ws, row, col, name, n_actions, detail_sheet, detail_row, is_init):
    r1, c1 = row, col
    r2, c2 = row + STEP_H_D - 1, col + STEP_W_D - 1
    fill = PatternFill('solid', start_color='0C4A6E') if is_init \
           else PatternFill('solid', start_color='1E40AF')
    try:
        ws.merge_cells(start_row=r1, start_column=c1, end_row=r2, end_column=c2)
    except Exception:
        pass
    cell = ws.cell(row=r1, column=c1)
    cell.value = '=HYPERLINK("#\'{}\'!A{}","{} ({} act)")'.format(
        detail_sheet, detail_row, name, n_actions)
    cell.font      = Font(name='Calibri', bold=True, size=8, color='FFFFFF')
    cell.fill      = fill
    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    cell.border    = Border(
        left=Side(style='medium', color='BFDBFE'),
        right=Side(style='medium', color='BFDBFE'),
        top=Side(style='medium', color='BFDBFE'),
        bottom=Side(style='medium', color='BFDBFE'),
    )
    return r1, c1, r2, c2

def _draw_trans_cell(ws, row, col, name, detail_sheet, detail_row, is_end):
    r1, c1 = row, col
    r2, c2 = row + TRAN_H_D - 1, col + TRAN_W_D - 1
    fill = PatternFill('solid', start_color='7F1D1D') if is_end \
           else PatternFill('solid', start_color='065F46')
    try:
        ws.merge_cells(start_row=r1, start_column=c1, end_row=r2, end_column=c2)
    except Exception:
        pass
    cell = ws.cell(row=r1, column=c1)
    cell.value = '=HYPERLINK("#\'{}\'!A{}","\u25c6 {}")'.format(
        detail_sheet, detail_row, name)
    cell.font      = Font(name='Calibri', bold=True, size=7, color='FFFFFF')
    cell.fill      = fill
    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    cell.border    = Border(
        left=Side(style='thin', color='6EE7B7'),
        right=Side(style='thin', color='6EE7B7'),
        top=Side(style='thin', color='6EE7B7'),
        bottom=Side(style='thin', color='6EE7B7'),
    )
    return r1, c1, r2, c2

def _draw_arrow(ws, from_row, to_row, mid_col, bg_fill):
    for r in range(from_row + 1, to_row):
        cell = ws.cell(row=r, column=mid_col)
        if cell.value is None:
            cell.fill  = bg_fill
            cell.value = ''

def build_sfc_diagram_sheet(wb, label, data, detail_name):
    """Build the _D diagram sheet for one logic block."""
    diag_name = (label[:27] + '_D')[:31]
    ws = wb.create_sheet(title=diag_name)

    steps = data['ordered_steps']
    trans = data['transitions']
    s2t   = data['step_to_trans']

    if not steps:
        ws['A1'] = '(No SFC steps to diagram)'
        return ws

    all_x = [v['x'] for _,v in steps] + [v['x'] for v in trans.values()]
    all_y = [v['y'] for _,v in steps] + [v['y'] for v in trans.values()]
    x_min, y_min = min(all_x), min(all_y)

    max_gc = max(round((x-x_min)*SCALE_D)+STEP_W_D+6 for x in all_x) + 5
    max_gr = max(round((y-y_min)*SCALE_D)+STEP_H_D+4 for y in all_y) + 5
    max_gc = max(max_gc, 30)

    ws.sheet_view.showGridLines = False
    for i in range(1, max_gc + 2):
        ws.column_dimensions[get_column_letter(i)].width = 8/7
    for i in range(1, max_gr + 4):
        ws.row_dimensions[i].height = 15

    # Light background
    bg = PatternFill('solid', start_color='F8FAFC')
    for r in range(1, max_gr + 3):
        for c in range(1, max_gc + 2):
            ws.cell(row=r, column=c).fill = bg

    # Title
    ws.row_dimensions[1].height = 22
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_gc)
    tc = ws.cell(row=1, column=1)
    tc.value     = '  SFC: {}   |   Click step (blue) or transition (green/red) for detail'.format(label)
    tc.font      = Font(name='Calibri', bold=True, size=10, color='FFFFFF')
    tc.fill      = PatternFill('solid', start_color='0F172A')
    tc.alignment = Alignment(horizontal='left', vertical='center')

    arrow_fill = PatternFill('solid', start_color='94A3B8')
    step_rows  = data.get('_detail_rows', {})
    trans_rows = data.get('_trans_detail_rows', {})

    step_boxes  = {}
    for si, (sname, sdata) in enumerate(steps):
        gr, gc = _to_grid(sdata['x'], sdata['y'], x_min, y_min)
        gr += 1  # title offset
        det = step_rows.get(sname, 2)
        r1,c1,r2,c2 = _draw_step_cell(ws, gr, gc, sname,
                                        len(sdata['actions']),
                                        detail_name, det, si == 0)
        step_boxes[sname] = (r1,c1,r2,c2)

    trans_boxes = {}
    for tname, tdata in trans.items():
        gr, gc = _to_grid(tdata['x'], tdata['y'], x_min, y_min)
        gr += 1
        det    = trans_rows.get(tname, 2)
        is_end = tdata.get('termination','F') == 'T'
        r1,c1,r2,c2 = _draw_trans_cell(ws, gr, gc, tname,
                                         detail_name, det, is_end)
        trans_boxes[tname] = (r1,c1,r2,c2)

    # Draw arrows step → transition
    for sname, tlist in s2t.items():
        if sname not in step_boxes: continue
        sr1,sc1,sr2,sc2 = step_boxes[sname]
        for tname in tlist:
            if tname not in trans_boxes: continue
            tr1,tc1_,tr2,tc2_ = trans_boxes[tname]
            mid = (tc1_ + tc2_) // 2
            if tr1 > sr2:
                _draw_arrow(ws, sr2, tr1, mid, arrow_fill)

    # Legend
    lr = max_gr + 2
    ws.row_dimensions[lr].height = 16
    ws.merge_cells(start_row=lr, start_column=1, end_row=lr, end_column=max_gc)
    lc = ws.cell(row=lr, column=1)
    lc.value     = '  Blue = Step  |  Green = Transition  |  Red = Terminating transition  |  Dark blue = Initial step  |  Click any shape to see detail'
    lc.font      = Font(name='Calibri', size=8, color='64748B', italic=True)
    lc.alignment = Alignment(horizontal='left', vertical='center')

    return ws

def build_detail_sheet_diag(wb, label, data, diag_name, first=False):
    """Build the _L detail sheet with back-links to the diagram."""
    detail_name = (label[:28] + '_L')[:31]
    ws = wb.active if first else wb.create_sheet(title=detail_name)
    if first: ws.title = detail_name

    NCOLS = 6
    for ci, w in enumerate([26,10,32,10,14,14], 1):
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.freeze_panes = 'A3'

    NAVY_F = PatternFill('solid', start_color='0F172A')
    BLUE_H = PatternFill('solid', start_color='1E3A8A')
    BLUE_S = PatternFill('solid', start_color='2563EB')
    GREEN_H= PatternFill('solid', start_color='065F46')
    ALT_   = PatternFill('solid', start_color='EFF6FF')
    ALT_G_ = PatternFill('solid', start_color='ECFDF5')
    WHITE_ = PatternFill('solid', start_color='FFFFFF')
    THIN__ = Side(style='thin', color='CBD5E1')
    BORD__ = Border(left=THIN__, right=THIN__, top=THIN__, bottom=THIN__)

    def sc(r, c, val='', bold=False, sz=10, fc='0F172A',
           fill=None, h='left', wrap=False, merge_to=None):
        cell = ws.cell(row=r, column=c,
                       value=str(val) if val is not None else '')
        cell.font      = Font(name='Calibri', bold=bold, size=sz, color=fc)
        cell.fill      = fill or WHITE_
        cell.alignment = Alignment(horizontal=h, vertical='top', wrap_text=wrap)
        cell.border    = BORD__
        if merge_to:
            ws.merge_cells(start_row=r, start_column=c,
                           end_row=r, end_column=merge_to)
        return cell

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=NCOLS)
    t = ws.cell(row=1, column=1,
        value='  Detail: {}   |   Click step or transition to return to diagram'.format(label))
    t.font = Font(name='Calibri', bold=True, size=12, color='FFFFFF')
    t.fill = NAVY_F
    t.alignment = Alignment(horizontal='left', vertical='center')
    ws.row_dimensions[1].height = 24

    for ci, h in enumerate(['Step / Transition','Action ID','Description',
                             'Qualifier','Expression','Delay / Confirm'], 1):
        sc(2, ci, h, bold=True, fc='FFFFFF', fill=BLUE_H, h='center')
    ws.row_dimensions[2].height = 16

    step_rows, trans_rows = {}, {}
    row = 3

    for si, (sname, sdata) in enumerate(data['ordered_steps']):
        step_rows[sname] = row
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=NCOLS)
        sh = ws.cell(row=row, column=1)
        sh.value = '=HYPERLINK("#\'{}\'!A1","  STEP {}:  {}   ({} actions)")'.format(
            diag_name, si+1, sname, len(sdata['actions']))
        sh.font  = Font(name='Calibri', bold=True, size=10, color='FFFFFF')
        sh.fill  = BLUE_S
        sh.alignment = Alignment(horizontal='left', vertical='center')
        sh.border= BORD__
        ws.row_dimensions[row].height = 17
        row += 1

        for ai, a in enumerate(sdata['actions']):
            f = ALT_ if ai % 2 == 0 else WHITE_
            delay   = a.get('delay_time','') or a.get('delay_expression','') or ''
            confirm = a.get('confirm_expression','') or ''
            extra   = ' / '.join(x for x in [delay, confirm] if x)
            sc(row,1, sname,             fill=f)
            sc(row,2, a['action'],       fill=f, h='center')
            sc(row,3, a['description'],  fill=f, wrap=True)
            sc(row,4, a['qualifier'],    fill=f, h='center')
            sc(row,5, a['expression'],   fill=f, wrap=True)
            sc(row,6, extra,             fill=f, wrap=True)
            ws.row_dimensions[row].height = max(15, min(60,
                15*max(1, len(a.get('expression',''))//55+1)))
            row += 1

        tlist = data['step_to_trans'].get(sname, [])
        for ti, tname in enumerate(tlist):
            tr = data['transitions'][tname]
            trans_rows[tname] = row
            term = 'END' if tr.get('termination') == 'T' else 'NEXT'
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=NCOLS)
            th = ws.cell(row=row, column=1)
            th.value = '=HYPERLINK("#\'{}\'!A1","  \u25c6 {}   ({})")'.format(
                diag_name, tname, term)
            th.font  = Font(name='Calibri', bold=True, size=9, color='FFFFFF')
            th.fill  = GREEN_H
            th.alignment = Alignment(horizontal='left', vertical='center')
            th.border= BORD__
            ws.row_dimensions[row].height = 15
            row += 1

            f = ALT_G_ if ti % 2 == 0 else WHITE_
            sc(row,1, tname,                   fill=f, bold=True)
            sc(row,2, term,                    fill=f, h='center')
            sc(row,3, tr.get('description',''),fill=f, wrap=True)
            sc(row,4, '',                      fill=f)
            ws.merge_cells(start_row=row, start_column=5, end_row=row, end_column=NCOLS)
            ec = ws.cell(row=row, column=5, value=tr.get('expression',''))
            ec.font      = Font(name='Calibri', size=9, italic=True, color='0F172A')
            ec.fill      = f
            ec.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            ec.border    = BORD__
            ws.row_dimensions[row].height = max(15, min(60,
                15*max(1, len(tr.get('expression',''))//80+1)))
            row += 1

        for ci in range(1, NCOLS+1):
            sc(row, ci, '', fill=WHITE_)
        ws.row_dimensions[row].height = 6
        row += 1

    return ws, detail_name, step_rows, trans_rows

def build_phase_diagram_excel(blocks, fname, opts):
    """Build an Excel workbook with SFC diagram + detail sheets per block."""
    wb   = openpyxl.Workbook()
    used = set()
    first = True

    for fb_name, data in blocks.items():
        lbl = derive_phase_label(fb_name, data.get('instance_name',''),
                                 data.get('description',''), used)
        diag_name = (lbl[:27] + '_D')[:31]

        # Detail sheet first (so we have row numbers for hyperlinks)
        _, detail_name, step_rows, trans_rows = build_detail_sheet_diag(
            wb, lbl, data, diag_name, first)
        first = False
        data['_detail_rows']       = step_rows
        data['_trans_detail_rows'] = trans_rows

        # Diagram sheet
        build_sfc_diagram_sheet(wb, lbl, data, detail_name)

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════════
# DDS WORD BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def build_dds_word(parsed_data, fhx_type, fname, opts):
    """
    Generate a pharma/biotech-style Design Description Specification (DDS)
    Word document from parsed FHX data.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Styles ────────────────────────────────────────────────────────────────
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)

    def add_heading(text, level=1, color='1F3864'):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor.from_string(color)
        p.runs[0].font.name = 'Calibri'
        return p

    def add_para(text='', bold=False, italic=False, size=10, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        return p

    def add_table_row(table, cells, bold=False, bg=None, font_size=9):
        row = table.add_row()
        for i, val in enumerate(cells):
            cell = row.cells[i]
            cell.text = str(val) if val else ''
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(font_size)
                    run.font.bold = bold
            if bg:
                _set_cell_bg(cell, bg)
        return row

    cols = opts.get('columns', ['step','description','action','qualifier',
                                'expression','delay','confirm_expression'])

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('DESIGN DESCRIPTION SPECIFICATION')
    run.font.name = 'Calibri'; run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string('1F3864')

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run(fname)
    run2.font.name = 'Calibri'; run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor.from_string('2E75B6')

    doc.add_paragraph()
    type_map = {'phase':'Phase Logic','em_cd':'Command Driven Equipment Module','em_sd':'State Driven Equipment Module'}
    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = type_p.add_run(type_map.get(fhx_type, fhx_type))
    run3.font.name = 'Calibri'; run3.font.size = Pt(12)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor.from_string('666666')

    doc.add_paragraph()

    # Cover table
    cover_tbl = doc.add_table(rows=0, cols=2)
    cover_tbl.style = 'Table Grid'
    for label, val in [
        ('Document Number', fname + '-DDS-001'),
        ('System / Module', fname),
        ('Document Type', 'Design Description Specification'),
        ('DeltaV Export Type', type_map.get(fhx_type, fhx_type)),
        ('Generated', datetime.datetime.now().strftime('%d-%b-%Y %H:%M')),
        ('Status', 'DRAFT'),
        ('Revision', '0.1'),
    ]:
        row = cover_tbl.add_row()
        row.cells[0].text = label
        row.cells[1].text = val
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True; run.font.name='Calibri'; run.font.size=Pt(10)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.name='Calibri'; run.font.size=Pt(10)
        _set_cell_bg(row.cells[0], 'D9E1F2')
    cover_tbl.columns[0].width = Cm(6)
    cover_tbl.columns[1].width = Cm(10)

    doc.add_page_break()

    # ── SECTION 1 — PURPOSE & SCOPE ───────────────────────────────────────────
    add_heading('1. Purpose and Scope', 1)
    add_para(
        f'This Design Description Specification (DDS) describes the control logic '
        f'for {fname}, as exported from the DeltaV Distributed Control System. '
        f'This document is intended to support validation activities (IQ/OQ/PQ), '
        f'change control review, and code review in accordance with 21 CFR Part 11 '
        f'and GAMP 5 guidelines.'
    )
    doc.add_paragraph()
    add_heading('1.1 Document Scope', 2)
    if fhx_type == 'phase':
        add_para('This document covers the Sequential Function Chart (SFC) logic for all phase '
                 'states including RUNNING, ABORTING, HOLDING, RESTARTING and STOPPING logic blocks.')
    elif fhx_type == 'em_cd':
        add_para('This document covers the Command Driven Equipment Module logic, describing '
                 'each command\'s sequential steps, actions, and transition conditions.')
    else:
        add_para('This document covers the State Driven Equipment Module logic, describing '
                 'each configured state and the corresponding device output positions.')

    doc.add_page_break()

    # ── SECTION 2 — REVISION HISTORY ─────────────────────────────────────────
    add_heading('2. Revision History', 1)
    rev_tbl = doc.add_table(rows=0, cols=4)
    rev_tbl.style = 'Table Grid'
    hrow = rev_tbl.add_row()
    for i, h in enumerate(['Rev', 'Date', 'Author', 'Description of Change']):
        hrow.cells[i].text = h
        _set_cell_bg(hrow.cells[i], '1F3864')
        for para in hrow.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold=True; run.font.name='Calibri'
                run.font.size=Pt(10); run.font.color.rgb=RGBColor.from_string('FFFFFF')
    drow = rev_tbl.add_row()
    drow.cells[0].text = '0.1'
    drow.cells[1].text = datetime.datetime.now().strftime('%d-%b-%Y')
    drow.cells[2].text = 'Auto-generated'
    drow.cells[3].text = 'Initial draft generated from DeltaV FHX export'
    for cell in drow.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name='Calibri'; run.font.size=Pt(10)

    doc.add_page_break()

    # ── SECTION 3 — OVERVIEW ──────────────────────────────────────────────────
    add_heading('3. Logic Overview', 1)

    if fhx_type == 'phase':
        blocks = parsed_data
        add_para(f'This phase contains {len(blocks)} logic block(s). '
                 f'Each block corresponds to a phase state or supporting function.')
        doc.add_paragraph()
        ov_tbl = doc.add_table(rows=0, cols=4)
        ov_tbl.style = 'Table Grid'
        hr = ov_tbl.add_row()
        for i, h in enumerate(['Logic Block','Description','Steps','Actions']):
            hr.cells[i].text = h
            _set_cell_bg(hr.cells[i], '1F3864')
            for para in hr.cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold=True; run.font.name='Calibri'
                    run.font.size=Pt(10); run.font.color.rgb=RGBColor.from_string('FFFFFF')
        for i, (fb, data) in enumerate(blocks.items()):
            acts = sum(len(s[1]['actions']) for s in data['ordered_steps'])
            lbl  = derive_phase_label(fb, data.get('instance_name',''),
                                      data.get('description',''), set())
            dr = ov_tbl.add_row()
            dr.cells[0].text = lbl
            dr.cells[1].text = data.get('description','')
            dr.cells[2].text = str(len(data['ordered_steps']))
            dr.cells[3].text = str(acts)
            bg = 'D9E1F2' if i%2==0 else 'FFFFFF'
            for cell in dr.cells:
                _set_cell_bg(cell, bg)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name='Calibri'; run.font.size=Pt(10)

    elif fhx_type == 'em_cd':
        commands = parsed_data
        add_para(f'This Command Driven EM contains {len(commands)} command(s).')
        doc.add_paragraph()
        ov_tbl = doc.add_table(rows=0, cols=4)
        ov_tbl.style = 'Table Grid'
        hr = ov_tbl.add_row()
        for i, h in enumerate(['Command','Description','Steps','Actions']):
            hr.cells[i].text = h
            _set_cell_bg(hr.cells[i], '1F3864')
            for para in hr.cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold=True; run.font.name='Calibri'
                    run.font.size=Pt(10); run.font.color.rgb=RGBColor.from_string('FFFFFF')
        for i, cmd in enumerate(commands):
            acts = sum(len(s[1]['actions']) for s in cmd['ordered_steps'])
            dr = ov_tbl.add_row()
            dr.cells[0].text = cmd['command_name']
            dr.cells[1].text = cmd.get('fb_description','')
            dr.cells[2].text = str(len(cmd['ordered_steps']))
            dr.cells[3].text = str(acts)
            bg = 'D9E1F2' if i%2==0 else 'FFFFFF'
            for cell in dr.cells:
                _set_cell_bg(cell, bg)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name='Calibri'; run.font.size=Pt(10)

    else:  # em_sd
        em_list = parsed_data
        for em in em_list:
            add_para(f'EM: {em["em_name"]} — {em["em_description"]}')
            add_para(f'Devices: {len(em["devices"])}   States: {len(em["states"])}')

    doc.add_page_break()

    # ── SECTION 4 — DETAILED LOGIC ────────────────────────────────────────────
    add_heading('4. Detailed Logic Description', 1)

    FIXED_COLS   = ['step', 'description', 'action']
    ALL_COLS     = ['step','description','action','qualifier',
                    'expression','delay','delay_expression',
                    'confirm_expression','confirm_timeout']
    COL_LABELS   = {
        'step': 'Step', 'description': 'Description', 'action': 'Action ID',
        'qualifier': 'Qualifier', 'expression': 'Expression',
        'delay': 'Delay', 'delay_expression': 'Delay Expression',
        'confirm_expression': 'Confirm Expression', 'confirm_timeout': 'Confirm Timeout'
    }
    # Only include selected columns (always keep fixed 3)
    sel_cols = FIXED_COLS + [c for c in cols if c not in FIXED_COLS and c in ALL_COLS]

    # Column widths in cm (total page width ~16cm)
    COL_W_MAP = {
        'step':4.0,'description':4.5,'action':2.5,'qualifier':1.8,
        'expression':5.5,'delay':1.8,'delay_expression':3.5,
        'confirm_expression':3.5,'confirm_timeout':2.0
    }
    # Scale widths to fit
    total_w = sum(COL_W_MAP[c] for c in sel_cols)
    scale   = 15.5 / total_w

    def make_logic_table():
        tbl = doc.add_table(rows=0, cols=len(sel_cols))
        tbl.style = 'Table Grid'
        hr  = tbl.add_row()
        for i, col in enumerate(sel_cols):
            hr.cells[i].text = COL_LABELS[col]
            _set_cell_bg(hr.cells[i], '1F3864')
            for para in hr.cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold=True; run.font.name='Calibri'
                    run.font.size=Pt(9); run.font.color.rgb=RGBColor.from_string('FFFFFF')
            hr.cells[i].width = Cm(COL_W_MAP[col] * scale)
        return tbl

    def add_step_header_row(tbl, step_num, step_name, action_count):
        row = tbl.add_row()
        for ci in range(len(sel_cols)):
            _set_cell_bg(row.cells[ci], '2E75B6')
        merged = row.cells[0]
        for ci in range(1, len(sel_cols)):
            merged = merged.merge(row.cells[ci])
        merged.text = f'  STEP {step_num}:  {step_name}   ({action_count} actions)'
        for para in merged.paragraphs:
            for run in para.runs:
                run.font.bold=True; run.font.name='Calibri'
                run.font.size=Pt(9); run.font.color.rgb=RGBColor.from_string('FFFFFF')

    def add_action_row(tbl, step_name, action, alt):
        row = tbl.add_row()
        bg  = 'D9E1F2' if alt else 'FFFFFF'
        vals = {
            'step': step_name,
            'description': action.get('description',''),
            'action': action.get('action',''),
            'qualifier': action.get('qualifier',''),
            'expression': action.get('expression',''),
            'delay': action.get('delay_time',''),
            'delay_expression': action.get('delay_expression',''),
            'confirm_expression': action.get('confirm_expression',''),
            'confirm_timeout': action.get('confirm_timeout',''),
        }
        for i, col in enumerate(sel_cols):
            row.cells[i].text = vals.get(col,'')
            _set_cell_bg(row.cells[i], bg)
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
            row.cells[i].width = Cm(COL_W_MAP[col] * scale)

    def add_transition_header(tbl, step_name, count):
        row = tbl.add_row()
        for ci in range(len(sel_cols)):
            _set_cell_bg(row.cells[ci], '4E7F2C')
        merged = row.cells[0]
        for ci in range(1, len(sel_cols)):
            merged = merged.merge(row.cells[ci])
        merged.text = f'  ↓  TRANSITIONS FROM  {step_name}  ({count})'
        for para in merged.paragraphs:
            for run in para.runs:
                run.font.bold=True; run.font.name='Calibri'
                run.font.size=Pt(8); run.font.color.rgb=RGBColor.from_string('FFFFFF')

    def add_transition_row(tbl, trans_name, trans, alt):
        row = tbl.add_row()
        bg  = 'E2EFDA' if alt else 'F2F9EE'
        vals = {
            'step': trans_name,
            'description': trans.get('description',''),
            'action': '→ NEXT' if trans.get('termination')!='T' else '⏹ END',
            'qualifier': '',
            'expression': trans.get('expression',''),
            'delay':'','delay_expression':'',
            'confirm_expression':'','confirm_timeout':'',
        }
        for i, col in enumerate(sel_cols):
            row.cells[i].text = vals.get(col,'')
            _set_cell_bg(row.cells[i], bg)
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
                    run.font.italic = (col=='expression')

    def write_sfc_section(data, section_title):
        add_heading(section_title, 2)
        tbl = make_logic_table()
        for si, (sn, sd) in enumerate(data['ordered_steps']):
            add_step_header_row(tbl, si+1, sn, len(sd['actions']))
            for ai, a in enumerate(sd['actions']):
                add_action_row(tbl, sn, a, ai%2==0)
            if opts.get('transitions', True):
                tl = data['step_to_trans'].get(sn, [])
                if tl:
                    add_transition_header(tbl, sn, len(tl))
                    for ti, tn in enumerate(tl):
                        add_transition_row(tbl, tn, data['transitions'][tn], ti%2==0)
        doc.add_paragraph()

    # Write detailed sections per type
    if fhx_type == 'phase':
        used = set()
        for fb, data in parsed_data.items():
            lbl = derive_phase_label(fb, data.get('instance_name',''),
                                     data.get('description',''), used)
            write_sfc_section(data, f'4.{list(parsed_data.keys()).index(fb)+1}  {lbl}')
            doc.add_paragraph()

    elif fhx_type == 'em_cd':
        for ci, cmd in enumerate(parsed_data):
            write_sfc_section(cmd, f'4.{ci+1}  Command: {cmd["command_name"]}')
            doc.add_paragraph()

    else:  # em_sd — state table
        for em in parsed_data:
            add_heading(f'4.1  {em["em_name"]} — State Table', 2)
            devices = em['devices']
            n_cols  = 3 + len(devices)
            sd_tbl  = doc.add_table(rows=0, cols=n_cols)
            sd_tbl.style = 'Table Grid'
            hr = sd_tbl.add_row()
            for i, h in enumerate(['#','State Name','Enabled'] + devices):
                hr.cells[i].text = h
                _set_cell_bg(hr.cells[i], '1F3864' if i < 3 else '7B3F00')
                for para in hr.cells[i].paragraphs:
                    for run in para.runs:
                        run.font.bold=True; run.font.name='Calibri'
                        run.font.size=Pt(9); run.font.color.rgb=RGBColor.from_string('FFFFFF')
            for ri, state in enumerate(em['states']):
                row = sd_tbl.add_row()
                row.cells[0].text = str(state['index'])
                row.cells[1].text = state['state_name']
                row.cells[2].text = 'Yes' if state['enabled'] else 'No'
                bg_base = 'F2F2F2' if ri%2==0 else 'FFFFFF'
                _set_cell_bg(row.cells[0], bg_base)
                _set_cell_bg(row.cells[1], bg_base)
                _set_cell_bg(row.cells[2], bg_base)
                for di, dev in enumerate(devices):
                    val = state['values'].get(dev,'')
                    dc  = state['dont_care'].get(dev, False)
                    row.cells[3+di].text = '—' if dc else val
                    bg = 'FFFFCC' if dc else ('C6EFCE' if val.upper()=='OPEN'
                         else 'FFCCCC' if val.upper()=='CLOSE' else bg_base)
                    _set_cell_bg(row.cells[3+di], bg)
                    for para in row.cells[3+di].paragraphs:
                        for run in para.runs:
                            run.font.name='Calibri'; run.font.size=Pt(9)
                for ci2 in range(3):
                    for para in row.cells[ci2].paragraphs:
                        for run in para.runs:
                            run.font.name='Calibri'; run.font.size=Pt(9)

    doc.add_page_break()

    # ── SECTION 5 — SIGN-OFF ──────────────────────────────────────────────────
    add_heading('5. Review and Approval', 1)
    add_para('This document requires review and approval before use in a GxP context.')
    doc.add_paragraph()
    sig_tbl = doc.add_table(rows=0, cols=4)
    sig_tbl.style = 'Table Grid'
    hr = sig_tbl.add_row()
    for i, h in enumerate(['Role','Name','Signature','Date']):
        hr.cells[i].text = h
        _set_cell_bg(hr.cells[i], '1F3864')
        for para in hr.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold=True; run.font.name='Calibri'
                run.font.size=Pt(10); run.font.color.rgb=RGBColor.from_string('FFFFFF')
    for role in ['Author','Technical Reviewer','Quality Reviewer','Approver']:
        row = sig_tbl.add_row()
        row.cells[0].text = role
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.name='Calibri'; run.font.size=Pt(10)
        for c in [1,2,3]:
            for para in row.cells[c].paragraphs:
                para.add_run('  ')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════════
# SHARED PARSE HELPER
# ═══════════════════════════════════════════════════════════════════════════════

def parse_and_build(text, fhx_type, fname, opts, output_format='excel'):
    """Parse text and build output. Returns (buf, sheet_names_or_none)."""
    if fhx_type == 'phase':
        data = parse_phase_fhx(text)
        if not data:
            raise ValueError('No phase logic blocks found. Is this a Phase FHX?')
        if output_format == 'word':
            return build_dds_word(data, fhx_type, fname, opts), None
        if output_format == 'diagram':
            return build_phase_diagram_excel(data, fname, opts), None
        buf = build_phase_excel(data, fname, opts)

    elif fhx_type == 'em_cd':
        data = parse_cdem_fhx(text)
        if not data:
            raise ValueError('No COMMAND_DRIVEN_ALGORITHM found.')
        if output_format == 'word':
            return build_dds_word(data, fhx_type, fname, opts), None
        buf = build_cdem_excel(data, fname, opts)

    elif fhx_type == 'em_sd':
        data = parse_sdem_fhx(text)
        if not data:
            raise ValueError('No STATE_DRIVEN_ALGORITHM found.')
        if output_format == 'word':
            return build_dds_word(data, fhx_type, fname, opts), None
        buf = build_sdem_excel(data, fname, opts)

    else:
        raise ValueError(f'Unknown type: {fhx_type}')

    # Get sheet names
    buf.seek(0)
    wb = openpyxl.load_workbook(buf, read_only=True)
    sheet_names = wb.sheetnames
    wb.close()
    buf.seek(0)
    return buf, sheet_names

# ═══════════════════════════════════════════════════════════════════════════════
# ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/')
def index():
    return send_file('index.html')

@app.route('/detect', methods=['POST'])
def detect():
    """Auto-detect the FHX type without full parsing."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file'}), 400
        f   = request.files['file']
        raw = f.read()
        text = decode_fhx(raw)

        detected = detect_fhx_type(text)

        # Fast stats using simple string counting (no regex backtracking)
        fb_count   = text.count('FUNCTION_BLOCK_DEFINITION NAME=')
        step_count = text.count('STEP NAME=')
        act_count  = text.count('ACTION NAME=')

        # Command names — only for small-enough files to avoid timeout
        cmd_names = []
        if len(text) < 500000:
            cmd_names = re.findall(r"A_COMMAND[^:]+:([^'\")\s]+)", text)
            cmd_names = [c.strip() for c in cmd_names if c.strip()][:20]

        return jsonify({
            'type':     detected,
            'fb_count': fb_count,
            'steps':    step_count,
            'actions':  act_count,
            'commands': cmd_names,
        })
    except Exception as e:
        # Always return JSON, never HTML
        return jsonify({'type': 'phase', 'fb_count': 0, 'steps': 0,
                        'actions': 0, 'commands': [], 'warning': str(e)}), 200

@app.route('/diagram')
def diagram_page():
    return send_file('diagram.html')

@app.route('/parse', methods=['POST'])
def parse_for_diagram():
    """Parse FHX and return structured JSON for the web diagram viewer."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file'}), 400
        f        = request.files['file']
        raw      = f.read()
        text     = decode_fhx(raw)
        fname    = re.sub(r'\.fhx$', '', f.filename, flags=re.IGNORECASE)
        fhx_type = detect_fhx_type(text)

        if fhx_type != 'phase':
            return jsonify({'error': 'SFC diagram is currently supported for Phase FHX only.'}), 400

        blocks = parse_phase_fhx(text)
        used   = set()
        result = {'filename': fname, 'type': fhx_type, 'blocks': {}}

        for fb_name, data in blocks.items():
            lbl = derive_phase_label(fb_name, data.get('instance_name',''),
                                     data.get('description',''), used)
            if not data['ordered_steps']:
                continue

            steps = []
            for sname, sdata in data['ordered_steps']:
                steps.append({
                    'name':        sname,
                    'x':           sdata['x'],
                    'y':           sdata['y'],
                    'actions':     sdata['actions'],
                    'transitions': data['step_to_trans'].get(sname, [])
                })

            trans = {}
            for tname, tdata in data['transitions'].items():
                trans[tname] = {
                    'name':        tname,
                    'x':           tdata['x'],
                    'y':           tdata['y'],
                    'description': tdata['description'],
                    'expression':  tdata['expression'],
                    'termination': tdata['termination']
                }

            result['blocks'][lbl] = {
                'label':       lbl,
                'description': data['description'],
                'steps':       steps,
                'transitions': trans
            }

        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/convert', methods=['POST'])
def convert():
    """Convert a single FHX file to Excel or Word."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    f = request.files['file']

    fhx_type = request.form.get('fhx_type', 'auto')
    fmt      = request.form.get('format', 'excel')   # excel | word
    cols_raw = request.form.get('columns', '')
    cols     = [c.strip() for c in cols_raw.split(',') if c.strip()] or None

    opts = {
        'summary':     request.form.get('summary',     'true') == 'true',
        'transitions': request.form.get('transitions', 'true') == 'true',
        'expressions': request.form.get('expressions', 'true') == 'true',
        'columns':     cols or ['step','description','action','qualifier',
                                'expression','delay','confirm_expression'],
    }

    raw   = f.read()
    text  = decode_fhx(raw)
    fname = re.sub(r'\.fhx$', '', f.filename, flags=re.IGNORECASE)

    if fhx_type == 'auto':
        fhx_type = detect_fhx_type(text)

    try:
        buf, sheet_names = parse_and_build(text, fhx_type, fname, opts, fmt)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': f'Parse error: {str(e)}'}), 500

    if fmt == 'word':
        resp = send_file(buf, as_attachment=True,
                         download_name=fname + '_DDS.docx',
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp.headers['X-Detected-Type'] = fhx_type
        return resp

    resp = send_file(buf, as_attachment=True,
                     download_name=fname + '_Logic.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    if sheet_names:
        resp.headers['X-Sheet-Names']   = '|'.join(sheet_names)
    resp.headers['X-Detected-Type'] = fhx_type
    return resp

@app.route('/batch', methods=['POST'])
def batch_convert():
    """
    Batch convert multiple FHX files into one combined Excel workbook.
    Each file gets its own tab group with a divider sheet.
    """
    import zipfile

    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'No files uploaded'}), 400

    fmt      = request.form.get('format', 'excel')
    cols_raw = request.form.get('columns', '')
    cols     = [c.strip() for c in cols_raw.split(',') if c.strip()] or None
    opts = {
        'summary':     request.form.get('summary',     'true') == 'true',
        'transitions': request.form.get('transitions', 'true') == 'true',
        'expressions': request.form.get('expressions', 'true') == 'true',
        'columns':     cols or ['step','description','action','qualifier',
                                'expression','delay','confirm_expression'],
    }

    if fmt == 'word':
        # Return zip of individual Word files
        zip_buf = io.BytesIO()
        results = []
        errors  = []
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for f in files:
                try:
                    raw      = f.read()
                    text     = decode_fhx(raw)
                    fname    = re.sub(r'\.fhx$', '', f.filename, flags=re.IGNORECASE)
                    fhx_type = detect_fhx_type(text)
                    buf, _   = parse_and_build(text, fhx_type, fname, opts, 'word')
                    zf.writestr(fname + '_DDS.docx', buf.read())
                    results.append(fname)
                except Exception as e:
                    errors.append(f'{f.filename}: {e}')
        zip_buf.seek(0)
        batch_name = 'DDS_Batch_' + datetime.datetime.now().strftime('%Y%m%d_%H%M') + '.zip'
        resp = send_file(zip_buf, as_attachment=True,
                         download_name=batch_name,
                         mimetype='application/zip')
        resp.headers['X-Batch-Count']  = str(len(results))
        resp.headers['X-Batch-Errors'] = str(len(errors))
        resp.headers['X-Batch-Names']  = '|'.join(results)
        return resp

    # Excel — combine into one workbook, one tab group per file
    combined_wb = openpyxl.Workbook()
    # Remove default empty sheet
    combined_wb.remove(combined_wb.active)

    batch_summary_rows = []
    errors = []
    file_count = 0

    for f in files:
        try:
            raw      = f.read()
            text     = decode_fhx(raw)
            fname    = re.sub(r'\.fhx$', '', f.filename, flags=re.IGNORECASE)
            fhx_type = detect_fhx_type(text)
            buf, sheet_names = parse_and_build(text, fhx_type, fname, opts, 'excel')

            # Load the individual workbook and copy sheets into combined
            buf.seek(0)
            src_wb = openpyxl.load_workbook(buf)

            # Add a divider sheet
            div_tab = (fname[:25] + '…') if len(fname) > 25 else fname
            div_ws  = combined_wb.create_sheet(title=('── ' + div_tab)[:31])
            div_ws.sheet_properties.tabColor = '1F3864'
            div_ws['A1'] = fname
            div_ws['A1'].font  = Font(name='Calibri', bold=True, size=14, color='FFFFFF')
            div_ws['A1'].fill  = PatternFill('solid', start_color='2C4A72')
            div_ws['A2'] = f'Type: {fhx_type.upper()}   |   Sheets: {len(src_wb.sheetnames)}'
            div_ws['A2'].font  = Font(name='Calibri', size=11, color='AAAAAA')
            div_ws.column_dimensions['A'].width = 60

            # Copy each sheet
            for sname in src_wb.sheetnames:
                src_ws = src_wb[sname]
                new_title = (fname[:15] + '·' + sname)[:31]
                new_ws = combined_wb.create_sheet(title=new_title)
                for row in src_ws.iter_rows():
                    for cell in row:
                        nc = new_ws.cell(row=cell.row, column=cell.column, value=cell.value)
                        if cell.has_style:
                            nc.font      = cell.font.copy()
                            nc.fill      = cell.fill.copy()
                            nc.border    = cell.border.copy()
                            nc.alignment = cell.alignment.copy()
                # Copy column widths
                for col_letter, cd in src_ws.column_dimensions.items():
                    new_ws.column_dimensions[col_letter].width = cd.width
                # Copy row heights
                for rn, rd in src_ws.row_dimensions.items():
                    new_ws.row_dimensions[rn].height = rd.height
                # Copy merges
                for merge in src_ws.merged_cells.ranges:
                    new_ws.merge_cells(str(merge))

            src_wb.close()
            acts = sum(len(re.findall(r'ACTION\s+NAME=', text)))
            batch_summary_rows.append({
                'name': fname, 'type': fhx_type,
                'sheets': len(sheet_names or []),
                'steps': len(re.findall(r'STEP\s+NAME=', text)),
                'actions': acts,
                'status': 'OK'
            })
            file_count += 1

        except Exception as e:
            errors.append({'name': f.filename, 'error': str(e)})
            batch_summary_rows.append({
                'name': f.filename, 'type': '?', 'sheets': 0,
                'steps': 0, 'actions': 0, 'status': f'ERROR: {e}'
            })

    # Add batch summary as first sheet
    if opts.get('summary', True):
        summ_ws = combined_wb.create_sheet(title='BATCH SUMMARY', index=0)
        summ_ws.sheet_properties.tabColor = '0D1B4B'
        summ_ws['A1'] = 'Batch Export Summary'
        summ_ws['A1'].font = Font(name='Calibri', bold=True, size=14, color='FFFFFF')
        summ_ws['A1'].fill = PatternFill('solid', start_color='1E3A5F')
        summ_ws['B1'] = datetime.datetime.now().strftime('%d-%b-%Y %H:%M')
        summ_ws['B1'].font = Font(name='Calibri', size=11, color='AAAAAA')
        summ_ws.merge_cells('A1:F1')

        headers = ['File', 'Type', 'Sheets', 'Steps', 'Actions', 'Status']
        for ci, h in enumerate(headers, 1):
            c = summ_ws.cell(row=2, column=ci, value=h)
            c.font = Font(name='Calibri', bold=True, color='FFFFFF', size=10)
            c.fill = PatternFill('solid', start_color='2C4A72')

        for ri, row_data in enumerate(batch_summary_rows):
            ri_sheet = ri + 3
            is_err   = row_data['status'] != 'OK'
            fill_col = 'FFCCCC' if is_err else ('D9E1F2' if ri%2==0 else 'FFFFFF')
            vals = [row_data['name'], row_data['type'].upper(),
                    row_data['sheets'], row_data['steps'],
                    row_data['actions'], row_data['status']]
            for ci, val in enumerate(vals, 1):
                c = summ_ws.cell(row=ri_sheet, column=ci, value=val)
                c.font = Font(name='Calibri', size=10)
                c.fill = PatternFill('solid', start_color=fill_col)

        col_widths = [40, 16, 8, 8, 10, 20]
        for ci, w in enumerate(col_widths, 1):
            summ_ws.column_dimensions[get_column_letter(ci)].width = w

    combined_buf = io.BytesIO()
    combined_wb.save(combined_buf)
    combined_buf.seek(0)

    batch_name = 'FHX_Batch_' + datetime.datetime.now().strftime('%Y%m%d_%H%M') + '.xlsx'
    resp = send_file(combined_buf, as_attachment=True,
                     download_name=batch_name,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp.headers['X-Batch-Count']  = str(file_count)
    resp.headers['X-Batch-Errors'] = str(len(errors))
    resp.headers['X-Batch-Names']  = '|'.join(r['name'] for r in batch_summary_rows)
    return resp

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
